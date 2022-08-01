# Import re for regex functions
import re
# Import docx to work with .docx files.
from docx import Document
# Import OS for navigating through the files in a directory
import os

# First, we create an empty lists to hold the words for and to replacement
replacement = []
to_replace = []

# Store file paths from input
directory = input("Enter the full path to the folder with the files you want to edit: ")
number_of_words = input('Write the number of words to replace: ')

for i in range(0, int(number_of_words)):
    to_replace.append(input(f'\n[{i + 1}] Write a word to replace: '))
    replacement.append(input('Write a replacement for this word: '))

for filename in os.listdir(directory):
    if filename.endswith(".docx"):
        print(f"\nCorrecting the file: {filename}")
        # full path to the current file
        filepath = os.path.join(directory, filename)
        document = Document(filepath)

        for i in range(0, int(number_of_words)):

            # Initialize the number of occurrences of this word to 0
            occurrences = {to_replace[i]: 0}

            # Loop through paragraphs
            for paragraph in document.paragraphs:
                # Loop through runs (style spans)
                for run in paragraph.runs:
                    # if there is text on this run, replace it
                    if run.text:
                        # get the replacement text
                        replaced_text = re.sub(to_replace[i], replacement[i], run.text, 999)
                        if replaced_text != run.text:
                            # if the replaced text is not the same as the original
                            # replace the text and increment the number of occurrences
                            run.text = replaced_text
                            occurrences[to_replace[i]] += 1

            # Loop through tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Loop through paragraphs
                        for paragraph in cell.paragraphs:
                            # Loop through runs (style spans)
                            for run in paragraph.runs:
                                # if there is text on this run, replace it
                                if run.text:
                                    # get the replacement text
                                    replaced_text = re.sub(to_replace[i], replacement[i], run.text, 999)
                                    if replaced_text != run.text:
                                        # if the replaced text is not the same as the original
                                        # replace the text and increment the number of occurrences
                                        run.text = replaced_text
                                        occurrences[to_replace[i]] += 1

            # print the number of occurrences of each word
            for word, count in occurrences.items():
                print(f"The word {word} was found and replaced by {replacement[i]} {count} time(s).")

            document.save(filepath)

input('\nPress any button to exit...')
