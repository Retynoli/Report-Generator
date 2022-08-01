# Import re for regex functions
import re
# Import sys for sys exit
from sys import exit
# Import docx to work with .docx files.
from docx import Document

# First, we create an empty lists to hold the path to all of lists files and words
list_paths = []
to_replace = []

# Store file paths from input
number_of_lists = input('Write the lists amount: ')
template_path = input('Write the full path to the template file: ')

for i in range(0, int(number_of_lists)):
    path = input(f'\n[{i + 1}] Write the full path to the list file: ')
    word = input('Write a word to replace for this list: ')

    if path.endswith('.docx') and template_path.endswith('.docx'):
        list_paths.append(path)
        to_replace.append(word)
    else:
        print('The file type is invalid, only .docx are supported')
        input('\nPress any button to exit...')
        exit()

counter = 0
words = Document(list_paths[0])

# Loop through replacer arguments
for replacement in words.paragraphs:
    template = Document(template_path)

    for i in range(0, int(number_of_lists)):
        words = Document(list_paths[i])
        cnt = 0
        # Loop through replacer arguments
        for replacement in words.paragraphs:
            # Find needed paragraphs
            if cnt != counter:
                cnt += 1
            else:
                print(f"\n[{counter + 1}] Сurrent replacement text: {replacement.text}")

                # Initialize the number of occurrences of this word to 0
                occurrences = {to_replace[i]: 0}

                # Loop through paragraphs
                for paragraph in template.paragraphs:
                    # Loop through runs (style spans)
                    for run in paragraph.runs:
                        # if there is text on this run, replace it
                        if run.text:
                            # get the replacement text
                            replaced_text = re.sub(to_replace[i], replacement.text, run.text, 999)
                            if replaced_text != run.text:
                                # if the replaced text is not the same as the original
                                # replace the text and increment the number of occurrences
                                run.text = replaced_text
                                occurrences[to_replace[i]] += 1

                # Loop through tables
                for table in template.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Loop through paragraphs
                            for paragraph in cell.paragraphs:
                                # Loop through runs (style spans)
                                for run in paragraph.runs:
                                    # if there is text on this run, replace it
                                    if run.text:
                                        # get the replacement text
                                        replaced_text = re.sub(to_replace[i], replacement.text, run.text, 999)
                                        if replaced_text != run.text:
                                            # if the replaced text is not the same as the original
                                            # replace the text and increment the number of occurrences
                                            run.text = replaced_text
                                            occurrences[to_replace[i]] += 1

                # Print the number of occurrences of each word
                for word, count in occurrences.items():
                    print(f"The word {word} was found and replaced by {replacement.text} {count} time(s).")

                break

    # make a new file name by changing the original file name with current word
    index = template_path.rfind('\\') + 1
    # replace all system reserved symbols if found
    filename = replacement.text.translate({ord(c): "-" for c in '\/:*?"<>|'})
    new_filepath = template_path[:index] + str(counter + 1) + ". " + filename + ".docx"

    print(f'File was saved at {new_filepath}')

    # save the new docx file
    template.save(new_filepath)
    counter += 1

input('\nPress any button to exit...')